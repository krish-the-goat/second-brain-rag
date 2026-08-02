import asyncio
import hashlib
from typing import List
import docx
from langchain_core.documents import Document

def _process_docx_sync(file_path: str) -> List[Document]:
    doc = docx.Document(file_path)
    documents = []
    current_section = "Unknown"
    
    for i, para in enumerate(doc.paragraphs):
        if i >= 10000:
            from app.core.exceptions import DocumentTooLargeError
            raise DocumentTooLargeError("DOCX exceeds the maximum allowed limit of 10,000 paragraphs. This is to prevent memory exhaustion (DoS).")
            
        if not para.text.strip():
            continue
            
        if para.style.name.startswith('Heading'):
            current_section = para.text.strip()
            
        text = para.text
        para_hash = hashlib.sha256(text.encode('utf-8')).hexdigest()
        
        metadata = {
            "filename": file_path.split("/")[-1],
            "section": current_section,
            "source_type": "docx",
            "hash": para_hash
        }
        
        documents.append(Document(page_content=text, metadata=metadata))
        
    return documents

async def load_docx(file_path: str) -> List[Document]:
    """Asynchronously load a docx and extract paragraphs with their sections."""
    return await asyncio.to_thread(_process_docx_sync, file_path)
