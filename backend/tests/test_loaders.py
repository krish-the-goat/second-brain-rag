import os
import pytest
from unittest.mock import patch, MagicMock
from langchain_core.documents import Document

from app.rag.loaders.pdf_loader import load_pdf
from app.rag.loaders.docx_loader import load_docx
from app.rag.loaders.web_loader import load_web
from app.rag.chunkers.recursive_chunker import chunk_documents
from app.core.exceptions import ProcessingError, ScrapingError

@pytest.fixture
def dummy_docx(tmp_path):
    import docx
    doc = docx.Document()
    doc.add_heading('Section 1', level=1)
    doc.add_paragraph('This is the first paragraph.')
    doc.add_paragraph('This is the second paragraph.')
    path = tmp_path / "dummy.docx"
    doc.save(str(path))
    return str(path)

@pytest.fixture
def dummy_pdf(tmp_path):
    from PyPDF2 import PdfWriter
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    writer.add_blank_page(width=72, height=72)
    path = tmp_path / "dummy.pdf"
    with open(path, "wb") as f:
        writer.write(f)
    return str(path)

@pytest.fixture
def dummy_url():
    return "https://example.com"

@pytest.mark.asyncio
@patch("PyPDF2.PdfReader")
@patch("pdfplumber.open")
async def test_pdf_loader_success(mock_plumber, mock_reader, dummy_pdf):
    # Mocking PyPDF2
    mock_pdf_instance = MagicMock()
    mock_pdf_instance.is_encrypted = False
    
    page1 = MagicMock()
    page1.extract_text.return_value = "Page 1 text"
    page2 = MagicMock()
    page2.extract_text.return_value = "Page 2 text"
    
    mock_pdf_instance.pages = [page1, page2]
    mock_reader.return_value = mock_pdf_instance
    
    # Mocking pdfplumber
    mock_plumber_instance = MagicMock()
    mock_plumber_page = MagicMock()
    mock_plumber_page.extract_tables.return_value = [["Header"], ["Row"]]
    mock_plumber_instance.__enter__.return_value.pages = [mock_plumber_page, mock_plumber_page]
    mock_plumber.return_value = mock_plumber_instance
    
    docs = await load_pdf(dummy_pdf)
    assert len(docs) == 2
    assert docs[0].metadata["page_number"] == 1
    assert docs[0].metadata["total_pages"] == 2
    assert docs[0].metadata["source_type"] == "pdf"
    assert "Page 1 text" in docs[0].page_content
    assert "Header" in docs[0].page_content

@pytest.mark.asyncio
@patch("PyPDF2.PdfReader")
async def test_pdf_loader_encrypted(mock_reader, dummy_pdf):
    mock_pdf_instance = MagicMock()
    mock_pdf_instance.is_encrypted = True
    mock_reader.return_value = mock_pdf_instance
    
    with pytest.raises(ProcessingError, match="Cannot process encrypted PDF."):
        await load_pdf(dummy_pdf)

@pytest.mark.asyncio
async def test_docx_loader_success(dummy_docx):
    docs = await load_docx(dummy_docx)
    assert len(docs) == 3
    assert docs[0].metadata["section"] == "Section 1"
    assert docs[0].metadata["source_type"] == "docx"
    assert "Section 1" in docs[0].page_content
    assert "This is the first paragraph." in docs[1].page_content
    assert "This is the second paragraph." in docs[2].page_content

@pytest.mark.asyncio
@patch("requests.get")
async def test_web_loader_success(mock_get, dummy_url):
    mock_response = MagicMock()
    mock_response.text = "<html><head><title>Test</title></head><body><nav>Nav</nav><p>Main content.</p><footer>Footer</footer></body></html>"
    mock_response.raise_for_status.return_value = None
    mock_get.return_value = mock_response
    
    docs = await load_web(dummy_url)
    assert len(docs) == 1
    assert "Main content." in docs[0].page_content
    assert "Nav" not in docs[0].page_content
    assert "Footer" not in docs[0].page_content
    assert docs[0].metadata["url"] == dummy_url
    assert docs[0].metadata["title"] == "Test"

@pytest.mark.asyncio
@patch("requests.get")
async def test_web_loader_scraping_error(mock_get, dummy_url):
    import requests
    mock_get.side_effect = requests.RequestException("Network error")
    
    with patch("time.sleep") as mock_sleep:
        with pytest.raises(ScrapingError):
            await load_web(dummy_url)
        assert mock_sleep.call_count == 3  # Exponential backoff 1, 2, 4

def test_recursive_chunker():
    doc = Document(page_content="A" * 3000, metadata={"source": "test"})
    
    os.environ["CHUNK_SIZE"] = "1000"
    os.environ["CHUNK_OVERLAP"] = "100"
    
    chunks = chunk_documents([doc])
    
    assert len(chunks) > 1
    for i, chunk in enumerate(chunks):
        assert chunk.metadata["source"] == "test"
        assert chunk.metadata["chunk_index"] == i
        assert chunk.metadata["total_chunks"] == len(chunks)
        assert len(chunk.page_content) <= 1000
