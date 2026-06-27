import os
from fpdf import FPDF
from docx import Document

os.makedirs('sample_docs', exist_ok=True)

# Document 1: Acme Corp Handbook (PDF)
pdf = FPDF()
pdf.add_page()
pdf.set_font("Arial", size=12)
text1 = """Acme Corp Employee Handbook - 2024 Edition

1. Introduction
Welcome to Acme Corp! Our core mission is to build the world's most robust anvils. Our headquarters is located at 123 Coyote Way, Desert City, NV. 
The current CEO is Jane Smith, who took over from John Doe in 2022.

2. Benefits
All full-time employees are eligible for the Gold Tier health plan. The Gold Tier plan covers 100% of dental and vision. It also includes a wellness stipend of $500 per year.
Vacation policy: Employees accrue 1.5 days of PTO per month. After 5 years of service, this increases to 2 days per month.

3. Code of Conduct
Employees must wear closed-toe shoes in the anvil manufacturing zone. 
Any safety violations should be reported directly to the Safety Coordinator, Mark Johnson. Mark's office is located in Building C, Room 402.
"""
# fpdf requires latin-1 encoded strings for basic font
for line in text1.split('\n'):
    pdf.multi_cell(0, 10, txt=line.encode('latin-1', 'replace').decode('latin-1'))
pdf.output("sample_docs/sample_1_acme_handbook.pdf")

# Document 2: Project Quantum Specs (PDF)
pdf = FPDF()
pdf.add_page()
pdf.set_font("Arial", size=12)
text2 = """Project Quantum: Technical Architecture

1. Overview
Project Quantum is our next-generation caching layer. It replaces the legacy Memcached cluster with a globally distributed Redis enterprise deployment.
The migration is scheduled for Q3 2024.

2. Performance Targets
The primary goal is to reduce P99 latency from 45ms to under 10ms for read operations. 
Write operations are expected to maintain a P95 latency of 15ms.

3. Incident Response
If the caching layer fails, the application will fallback to direct database queries. The database cluster is managed by the DBA team led by Sarah Connor. 
If database load exceeds 80%, circuit breakers will automatically shed non-critical traffic.
"""
for line in text2.split('\n'):
    pdf.multi_cell(0, 10, txt=line.encode('latin-1', 'replace').decode('latin-1'))
pdf.output("sample_docs/sample_2_quantum_specs.pdf")

# Document 3: Q4 Financial Report (DOCX)
doc = Document()
doc.add_heading('Q4 Financial Summary - Acme Corp', 0)
doc.add_paragraph('Q4 2023 was a record-breaking quarter for Acme Corp. Total revenue reached $45 million, representing a 20% year-over-year growth.')
doc.add_heading('Revenue Breakdown', level=1)
doc.add_paragraph('The standard anvil line contributed $25 million. The newly introduced rocket-powered skates generated $15 million. Services and maintenance made up the remaining $5 million.')
doc.add_heading('Operating Expenses', level=1)
doc.add_paragraph('Operating expenses increased by 15% due to aggressive R&D investments in the skate division and the hiring of 50 new engineers.')
doc.add_heading('Future Outlook', level=1)
doc.add_paragraph('For Q1 2024, the company projects revenue between $48M and $50M, assuming the supply chain issues with rocket fuel are resolved by February.')
doc.save('sample_docs/sample_3_financial_report.docx')

print("Generated sample documents successfully!")
